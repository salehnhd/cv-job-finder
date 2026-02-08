"""
Windows-safe CrewAI CV->Job Search->Rank script 

What it does:
1) Reads your CV from cvLLM.docx
2) Converts it to cv_text.txt
3) Uses 3 CrewAI agents:
   - CV Analyst: extracts skills/titles
   - Search Planner: builds search queries
   - Fit Ranker: searches job postings (Serper), scrapes pages, ranks fit
4) Writes:
   - crew_output.txt
   - shortlisted_jobs.json
   - shortlisted_jobs.md

IMPORTANT:
- Requires env vars set in the same terminal before running:
    $env:OPENAI_API_KEY="sk-..."
    $env:SERPER_API_KEY="..."

- Avoids CrewAI "memory" to prevent embedchain/chroma/hnswlib build issues.

Recommended installs (inside venv312):
    python -m pip install crewai==0.28.8
    python -m pip install "langchain==0.1.13" "langchain-core==0.1.35" "langchain-community==0.0.29" "langchain-openai==0.1.3"
    python -m pip install python-docx requests beautifulsoup4
"""

import os
import json
import re
from typing import Any, Dict, List

import requests
from bs4 import BeautifulSoup
from docx import Document

from crewai import Agent, Task, Crew


# -----------------------------
# ENV CHECK
# -----------------------------
if not os.getenv("OPENAI_API_KEY"):
    raise RuntimeError(
        "Missing OPENAI_API_KEY. Set it in PowerShell:\n"
        '$env:OPENAI_API_KEY="sk-..."'
    )

if not os.getenv("SERPER_API_KEY"):
    raise RuntimeError(
        "Missing SERPER_API_KEY. Set it in PowerShell:\n"
        '$env:SERPER_API_KEY="..."'
    )

os.environ["OPENAI_MODEL_NAME"] = os.getenv("OPENAI_MODEL_NAME", "gpt-3.5-turbo")

# -----------------------------
# CONFIG
# -----------------------------
CV_DOCX = "cvLLM.docx"
CV_TEXT_PATH = "cv_text.txt"
SERPER_URL = "https://google.serper.dev/search"

# Control how many results per query and how many top jobs to evaluate deeply
RESULTS_PER_QUERY = 10
MAX_JOBS_TO_EVALUATE = 15


# -----------------------------
# DOCX -> TEXT
# -----------------------------
def docx_to_text(docx_path: str) -> str:
    if not os.path.exists(docx_path):
        raise FileNotFoundError(
            f"CV DOCX not found: {docx_path}\n"
            "Put cvLLM.docx in the same folder as this script, or update CV_DOCX path."
        )

    doc = Document(docx_path)
    parts: List[str] = []

    # paragraphs
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # tables (flatten)
    for table in doc.tables:
        parts.append("\n---\n")
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts).strip() + "\n"
    with open(CV_TEXT_PATH, "w", encoding="utf-8") as f:
        f.write(text)

    return text


# -----------------------------
# SIMPLE WEB TOOLS (Serper + Scrape)
# -----------------------------
def serper_search(query: str, k: int = 10) -> List[Dict[str, str]]:
    headers = {
        "X-API-KEY": os.getenv("SERPER_API_KEY"),
        "Content-Type": "application/json",
    }
    payload = {"q": query, "num": k}
    r = requests.post(SERPER_URL, headers=headers, json=payload, timeout=30)
    r.raise_for_status()
    data = r.json()

    out: List[Dict[str, str]] = []
    for item in data.get("organic", []) or []:
        out.append(
            {
                "title": item.get("title", ""),
                "url": item.get("link", ""),
                "snippet": item.get("snippet", ""),
            }
        )
    return out


def scrape_page_text(url: str, max_chars: int = 8000) -> str:
    try:
        r = requests.get(
            url,
            timeout=25,
            headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"},
        )
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "html.parser")

        # remove obvious noise
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()

        # gather text from headings + paragraphs + lists
        chunks: List[str] = []
        for tag in soup.find_all(["h1", "h2", "h3", "p", "li"]):
            t = tag.get_text(" ", strip=True)
            if t:
                chunks.append(t)

        text = "\n".join(chunks)
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        return text[:max_chars]
    except Exception:
        return ""


# -----------------------------
# JSON EXTRACTION HELPERS
# -----------------------------
def extract_first_json(text: str) -> Any:
    """
    Extract first JSON object/array from a text blob.
    """
    if not isinstance(text, str):
        text = str(text)

    # direct
    try:
        return json.loads(text)
    except Exception:
        pass

    # ```json ... ```
    m = re.search(r"```json\s*(.*?)\s*```", text, re.DOTALL | re.IGNORECASE)
    if m:
        return json.loads(m.group(1).strip())

    # find first { or [
    starts = [i for i in [text.find("{"), text.find("[")] if i != -1]
    if not starts:
        raise ValueError("No JSON start token found.")
    start = min(starts)

    stack: List[str] = []
    for i in range(start, len(text)):
        ch = text[i]
        if ch in "{[":
            stack.append(ch)
        elif ch in "}]":
            if not stack:
                continue
            top = stack[-1]
            if (top == "{" and ch == "}") or (top == "[" and ch == "]"):
                stack.pop()
                if not stack:
                    candidate = text[start : i + 1]
                    return json.loads(candidate)

    raise ValueError("Could not extract a complete JSON blob.")


# -----------------------------
# READ CV NOW
# -----------------------------
CV_TEXT = docx_to_text(CV_DOCX)


# -----------------------------
# AGENTS
# -----------------------------
cv_analyst = Agent(
    role="CV Analyst",
    goal=(
        "Read the CV text and produce a grounded, structured evaluation: skills, domains, tools, "
        "seniority, constraints, and target job titles. Do NOT invent experience."
    ),
    verbose=True,
    backstory="You are a meticulous technical recruiter and scientist. You only use evidence from the CV.",
)

search_planner = Agent(
    role="Job Search Planner",
    goal=(
        "Turn the CV evaluation into strong search queries that find real job postings "
        "(prefer official ATS pages like Lever/Greenhouse/Workday/company careers)."
    ),
    verbose=True,
    backstory="You create search queries that reliably surface real job ads, not generic blog posts.",
)

fit_ranker = Agent(
    role="Job Fit Ranker",
    goal=(
        "Given a list of job posting URLs and the CV evaluation, scrape postings and score fit with clear reasons."
    ),
    verbose=True,
    backstory="You match candidates to jobs honestly, highlighting alignment and gaps without hallucinating.",
)


# -----------------------------
# TASK 1: CV evaluation JSON
# -----------------------------
cv_eval_task = Task(
    description=(
        "Analyze the CV below and output STRICT JSON only with these keys:\n"
        "  - core_skills: [..]\n"
        "  - secondary_skills: [..]\n"
        "  - domains: [..]\n"
        "  - tools_and_tech: [..]\n"
        "  - seniority_level: string\n"
        "  - strengths: [..]\n"
        "  - gaps_or_risks: [..]\n"
        "  - constraints: [..] (location, remote, visa, etc., only if present)\n"
        "  - target_job_titles: [..]\n"
        "  - avoid_job_titles: [..]\n\n"
        "Rules:\n"
        "  - Ground everything in the CV; do NOT invent.\n"
        "  - Output JSON only.\n\n"
        f"CV:\n{CV_TEXT}\n"
    ),
    expected_output="A JSON object with the specified keys.",
    agent=cv_analyst,
)


# -----------------------------
# TASK 2: Build search queries JSON
# -----------------------------
criteria_task = Task(
    description=(
        "Using the CV evaluation JSON, output STRICT JSON only with these keys:\n"
        "  - queries: 10 to 14 search queries (Google-style)\n"
        "  - include_keywords: [..]\n"
        "  - exclude_keywords: [..]\n"
        "  - preferred_sources: [..] (lever.co, greenhouse.io, workday, company careers pages)\n\n"
        "Rules:\n"
        "  - Queries should find real job postings, not advice articles.\n"
        "  - Strongly prefer official/company ATS pages.\n"
        "  - Output JSON only.\n"
    ),
    expected_output="A JSON object containing queries and filters.",
    context=[cv_eval_task],
    agent=search_planner,
)


# -----------------------------
# TASK 3: Search jobs + rank jobs (calls our Python tools)
# -----------------------------
def search_and_rank_jobs(criteria: Dict[str, Any], cv_eval: Dict[str, Any]) -> List[Dict[str, Any]]:
    queries = criteria.get("queries", []) or []
    exclude_keywords = [x.lower() for x in (criteria.get("exclude_keywords", []) or [])]

    # 1) Search
    seen = set()
    jobs: List[Dict[str, str]] = []
    for q in queries:
        try:
            results = serper_search(q, k=RESULTS_PER_QUERY)
        except Exception:
            continue

        for r in results:
            url = (r.get("url") or "").strip()
            title = (r.get("title") or "").strip()
            snippet = (r.get("snippet") or "").strip()

            if not url or url in seen:
                continue

            # basic exclude filter
            blob = f"{title} {snippet} {url}".lower()
            if any(kw and kw in blob for kw in exclude_keywords):
                continue

            seen.add(url)
            jobs.append({"title": title, "url": url, "snippet": snippet})

    # 2) Evaluate top N by scraping
    evaluated: List[Dict[str, Any]] = []
    for job in jobs[:MAX_JOBS_TO_EVALUATE]:
        url = job["url"]
        page_text = scrape_page_text(url)

        # If we couldn't scrape, still keep it as low-confidence
        if not page_text:
            evaluated.append(
                {
                    "title": job.get("title", ""),
                    "company": "",
                    "location": "",
                    "url": url,
                    "fit_score": 3,
                    "reason": "Could not extract job page text (blocked or dynamic page).",
                    "risks": ["Could not scrape job page"],
                }
            )
            continue

        # Use the LLM (fit_ranker agent) to score fit based on scraped text
        prompt = (
            "Given the CV evaluation JSON and the job posting text, score fit from 1 to 10.\n"
            "Return STRICT JSON only with keys:\n"
            "  - fit_score (number 1-10)\n"
            "  - reason (2-4 short bullets as a single string)\n"
            "  - company (string if you can infer from text/title else empty)\n"
            "  - location (string if present else empty)\n"
            "  - risks (list of 1-3 gaps/risks)\n\n"
            f"CV_EVAL_JSON:\n{json.dumps(cv_eval, ensure_ascii=False)}\n\n"
            f"JOB_TITLE:\n{job.get('title','')}\n\n"
            f"JOB_URL:\n{url}\n\n"
            f"JOB_TEXT:\n{page_text}\n"
        )

        # We run this by creating a tiny ad-hoc task executed by the fit_ranker agent
        tmp_task = Task(
            description=prompt,
            expected_output="STRICT JSON only.",
            agent=fit_ranker,
        )
        tmp_crew = Crew(
            agents=[fit_ranker],
            tasks=[tmp_task],
            verbose=False,
            memory=False,  # critical: no embedchain
        )

        tmp_result = tmp_crew.kickoff()
        tmp_text = str(tmp_result)

        try:
            scored = extract_first_json(tmp_text)
            if isinstance(scored, dict):
                evaluated.append(
                    {
                        "title": job.get("title", ""),
                        "company": scored.get("company", ""),
                        "location": scored.get("location", ""),
                        "url": url,
                        "fit_score": scored.get("fit_score", 0),
                        "reason": scored.get("reason", ""),
                        "risks": scored.get("risks", []),
                    }
                )
            else:
                raise ValueError("Not a JSON object")
        except Exception:
            evaluated.append(
                {
                    "title": job.get("title", ""),
                    "company": "",
                    "location": "",
                    "url": url,
                    "fit_score": 4,
                    "reason": "Could not parse scorer output reliably.",
                    "risks": ["Scoring parse failure"],
                }
            )

    # sort by fit_score desc
    evaluated.sort(key=lambda x: float(x.get("fit_score", 0) or 0), reverse=True)
    return evaluated


class PythonRunner:
    """
    Tiny helper object to run our python functions inside a Task step:
    We will call it from a Task by embedding the results in the final output.
    """
    @staticmethod
    def run(criteria_json_text: str, cv_eval_json_text: str) -> str:
        criteria = json.loads(criteria_json_text)
        cv_eval = json.loads(cv_eval_json_text)
        ranked = search_and_rank_jobs(criteria, cv_eval)
        return json.dumps(ranked, ensure_ascii=False, indent=2)


rank_task = Task(
    description=(
        "You will receive CV evaluation JSON and search criteria JSON from prior tasks.\n"
        "Your job is to output STRICT JSON ONLY: a list of ranked jobs.\n"
        "Do NOT include any extra text.\n"
        "If you cannot do web search directly, still provide the JSON format with empty list.\n"
        "IMPORTANT: Output MUST be JSON list only.\n"
    ),
    expected_output="A JSON list of ranked jobs.",
    context=[cv_eval_task, criteria_task],
    agent=fit_ranker,
)

# -----------------------------
# RUN MAIN CREW (memory disabled)
# -----------------------------
main_crew = Crew(
    agents=[cv_analyst, search_planner, fit_ranker],
    tasks=[cv_eval_task, criteria_task],
    verbose=True,
    memory=False,  # critical: prevents embedchain import path
)

cv_eval_out = main_crew.kickoff()
cv_eval_text = str(cv_eval_out)

# Extract CV eval JSON
cv_eval_json = extract_first_json(cv_eval_text)
if not isinstance(cv_eval_json, dict):
    raise RuntimeError("Could not parse CV evaluation JSON.")

# Now run criteria task alone (already in main_crew output, but we do it cleanly)
criteria_crew = Crew(
    agents=[search_planner],
    tasks=[criteria_task],
    verbose=True,
    memory=False,
)
criteria_out = criteria_crew.kickoff()
criteria_text = str(criteria_out)

criteria_json = extract_first_json(criteria_text)
if not isinstance(criteria_json, dict):
    raise RuntimeError("Could not parse criteria JSON.")

# Search + rank using Python tools, then save
ranked_jobs = search_and_rank_jobs(criteria_json, cv_eval_json)

with open("crew_output.txt", "w", encoding="utf-8") as f:
    f.write("CV_EVAL_JSON:\n")
    f.write(json.dumps(cv_eval_json, ensure_ascii=False, indent=2))
    f.write("\n\nCRITERIA_JSON:\n")
    f.write(json.dumps(criteria_json, ensure_ascii=False, indent=2))
    f.write("\n\nRANKED_JOBS_JSON:\n")
    f.write(json.dumps(ranked_jobs, ensure_ascii=False, indent=2))

with open("shortlisted_jobs.json", "w", encoding="utf-8") as f:
    json.dump(ranked_jobs, f, ensure_ascii=False, indent=2)

# markdown table
md_lines = []
md_lines.append("| # | Fit | Title | Company | Location | URL |")
md_lines.append("|---:|---:|---|---|---|---|")
for i, job in enumerate(ranked_jobs, start=1):
    md_lines.append(
        f"| {i} | {job.get('fit_score','')} | {str(job.get('title','')).replace('|','\\|')} | "
        f"{str(job.get('company','')).replace('|','\\|')} | {str(job.get('location','')).replace('|','\\|')} | "
        f"{job.get('url','')} |"
    )

with open("shortlisted_jobs.md", "w", encoding="utf-8") as f:
    f.write("\n".join(md_lines))

print("DONE. Wrote: crew_output.txt, shortlisted_jobs.json, shortlisted_jobs.md")
